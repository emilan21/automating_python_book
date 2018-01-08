#!/usr/bin/python3.5

# invitations.py - Program generates a word document with custom invitations from the input of a text file.

import docx
import sys

filename = sys.argv[1]

doc = docx.Document('invitations.docx')
p_count = 0

text_file = open(filename, 'r') 
text_file_content = text_file.readlines()
for string in text_file_content:
    p_count += 5
    doc.add_paragraph('It would be a pleasure to have the company of', style='italic_center')
    doc.add_paragraph(string, style='name_bold')
    doc.add_paragraph('at 11010 Memory Lane on the Evening of', style='italic_center')
    doc.add_paragraph('April 1st', 'normal_center')
    doc.add_paragraph("at 7 0'clock", style='italic_center')
    doc.paragraphs[p_count].runs[0].add_break(docx.text.run.WD_BREAK.PAGE)
text_file.close()

doc.save('new_invitations.docx')